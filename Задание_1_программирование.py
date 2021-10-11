import pandas as pd
import numpy as np
import math
import docx
import os
# Чтение файла в структуру данных Pandas DataFrame
data = pd.read_csv('Тестовое задание 1.csv', sep = ';')

# Выборка тех строк, где оценивается параметр 1. Такое решение основано на том факте, что нужно составить отчет по
# Summary of Efficacy Parametr 1.
data_new = data.query('PARAMCD=="EFF01"').iloc[0:]

# Удаление тех строк, где присутсвтуют незапланированные визиты. Возможно, если бы незапланированных визитов было больше,
# можно было бы составить аналитику и по ним, но так как он всего один, составить какую-либо статистику не предоставляется
# возможным. Также мной были выбраны те строки, где значения флага ITT равно 1, так как аналитика составляется по
# Intention-to-Treat population.
data_new = data_new.loc[(data_new['VISIT'] != 'Unscheduled') & (data_new['ITTFL'] == '1')]

# Перевод значений стобцов "VISITN" и "AVAL" в числовые значения для дальнейшей работы с ними.
data_new[["VISITN", "AVAL"]] = data_new[["VISITN", "AVAL"]].apply(pd.to_numeric)

# Замена пропущенных значений NaN на среднее значение по той же группе и визиту.
for row in data_new.itertuples(index=True):
    if math.isnan(row[6]):
        param = 'VISITN==' + str(row[2]) + ' and TRTGRPN==' + '"' + row[7] + '"'
        data_new.loc[row[0], 'AVAL'] = round(data_new.query(param)['AVAL'].mean(),1)

# Разделение данных по Treatment group для удобства работы с данными.
data_group_1 = data_new.query('TRTGRPN=="1"')
data_group_2 = data_new.query('TRTGRPN=="2"')

# Статистические параметры.
agg_func_math = {
    'AVAL': ['mean', 'min', 'max', 'std'],
    'SUBJID':['count']
}

# Данные для отчета, то есть результат вычисляемых статистических значений с номером группы и номером визита, я решил
# хранить в словаре, так как это удобно, если, например, нужно передать эти данные. Для данной задачи можно использовать
# текстовый формат JSON.
results = {}

# Результаты вычисленных статистических значений для двух групп, сгруппированные по номеру визита.
# Представление в виде таблиц.
group_1_stat = data_group_1.groupby(['VISITN']).agg(agg_func_math)
group_2_stat = data_group_2.groupby(['VISITN']).agg(agg_func_math)

# Текущий результат. Данный словарь используется для динамического добавления данных из таблиц "group_1_stat" и
# "group_2_stat" в результирующий словарь "results"
current_result = {'Treatment group 1':{'n':'', 'Mean':'', 'Standard':'', 'Minimum':'', 'Maximum':''},
                       'Treatment group 2':{'n':'', 'Mean':'', 'Standard':'', 'Minimum':'', 'Maximum':''}}

for i in range(1, max(len(group_1_stat), len(group_2_stat)) + 1):
    current_result['Treatment group 1']['n'] = round(group_1_stat['SUBJID']['count'][i])
    current_result['Treatment group 1']['Mean'] = round(group_1_stat['AVAL']['mean'][i], 2)
    current_result['Treatment group 1']['Standard'] = round(group_1_stat['AVAL']['std'][i], 2)
    current_result['Treatment group 1']['Minimum'] = round(group_1_stat['AVAL']['min'][i])
    current_result['Treatment group 1']['Maximum'] = round(group_1_stat['AVAL']['max'][i])

    current_result['Treatment group 2']['n'] = round(group_2_stat['SUBJID']['count'][i])
    current_result['Treatment group 2']['Mean'] = round(group_2_stat['AVAL']['mean'][i], 2)
    current_result['Treatment group 2']['Standard'] = round(group_2_stat['AVAL']['std'][i], 2)
    current_result['Treatment group 2']['Minimum'] = round(group_2_stat['AVAL']['min'][i])
    current_result['Treatment group 2']['Maximum'] = round(group_2_stat['AVAL']['max'][i])

    results[i] = current_result
    current_result = {'Treatment group 1': {'n': '', 'Mean': '', 'Standard': '', 'Minimum': '', 'Maximum': ''},
                      'Treatment group 2': {'n': '', 'Mean': '', 'Standard': '', 'Minimum': '', 'Maximum': ''}}

# Составления таблицы в формате docx. Для составления таблицы используются данные из словаря results.
doc = docx.Document()

doc.add_heading('Table: Summary of Efficacy Parameter 1 by Visit. Intention-to-Treat population.',2)

menuTable = doc.add_table(rows=1,cols=4)
menuTable.style= 'Table Grid'
hdr_Cells = menuTable.rows[0].cells
hdr_Cells[0].text = 'Visit'
hdr_Cells[1].text = 'Statistics'
hdr_Cells[2].text = 'Treatment group 1 (N=' + str(len(data_group_1['SUBJID'].unique())) + ')'
hdr_Cells[3].text = 'Treatment group 2 (N=' + str(len(data_group_2['SUBJID'].unique())) + ')'

records = []

for visit in results:
    current_list = []
    flag = True
    for statistics in results[visit]['Treatment group 1']:
        if flag:
            current_list.append('Visit ' + str(visit))
            flag = False
        else:
            current_list.append('')
        current_list.append(statistics)
        current_list.append(results[visit]['Treatment group 1'][statistics])
        current_list.append(results[visit]['Treatment group 2'][statistics])
        records.append(current_list)
        current_list = []

for VISIT, Statistics, group1, group2 in records:
    row_Cells = menuTable.add_row().cells
    row_Cells[0].text= str(VISIT)
    row_Cells[1].text = str(Statistics)
    row_Cells[2].text = str(group1)
    row_Cells[3].text = str(group2)

doc.save('results.docx')